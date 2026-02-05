# modules/diff_report.py
import streamlit as st
from docx import Document
import io


def diff_report_block():
    st.subheader("📄 差异分析结果报告导出")

    diff_keys = [
        k.replace("diff_result_", "")
        for k in st.session_state.keys()
        if k.startswith("diff_result_")
    ]

    if not diff_keys:
        st.info("暂无差异分析结果")
        return

    selected = st.selectbox("选择要导出的差异分析结果", diff_keys)

    if st.button("📄 生成 Word 报告"):
        res = st.session_state[f"diff_result_{selected}"]
        sig = st.session_state[f"sig_genes_{selected}"]

        doc = Document()
        doc.add_heading("差异分析报告", level=1)
        doc.add_paragraph(f"分析组合：{selected}")
        doc.add_paragraph(f"显著基因数：{len(sig)}")

        doc.add_heading("显著基因列表（前 50）", level=2)
        table = doc.add_table(rows=1, cols=len(sig.columns) + 1)
        hdrs = ["Gene"] + list(sig.columns)
        for i, h in enumerate(hdrs):
            table.rows[0].cells[i].text = h

        for gene, row in sig.head(50).iterrows():
            cells = table.add_row().cells
            cells[0].text = gene
            for i, v in enumerate(row.values):
                cells[i + 1].text = str(v)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        st.download_button(
            "📥 下载 Word 报告",
            buf,
            file_name=f"diff_report_{selected}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
